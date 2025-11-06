from flask import Flask, request, jsonify
import tempfile
import shutil
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from nltk.corpus import stopwords  # Day 4: Stopwords ke liye
from flask_cors import CORS
import re  # Day 4: Regular Expressions (Cleaning) ke liye
import os  # For file extension check
import io  # For in-memory image handling
from collections import Counter
import math

# PDF and DOCX support
try: 
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. PDF support disabled. Run: pip install PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled. Run: pip install python-docx")

# OCR support for image-based PDFs
try:
    from pdf2image import convert_from_bytes
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
    # Set Tesseract path if needed (Windows example)
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR not installed. Image-based PDF support disabled. Run: pip install pdf2image pytesseract pillow")

# Flask app ko initialize karein
app = Flask(__name__)
CORS(app)

# set up basic logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure NLTK stopwords corpus is available at startup (download if missing)
try:
    # This will raise LookupError if the corpus isn't available
    _ = stopwords.words('english')
except LookupError:
    try:
        import nltk
        print("NLTK stopwords not found. Downloading 'stopwords' corpus...")
        nltk.download('stopwords')
    except Exception as _e:
        print(f"Failed to download NLTK stopwords: {_e}")

# Use a module-level STOPWORDS so we don't attempt to load it repeatedly
try:
    STOPWORDS = set(stopwords.words('english'))
except Exception:
    # Fallback to a small built-in stop list if something goes wrong
    STOPWORDS = {
        'a','an','and','the','is','in','it','of','to','for','on','with','as','by'
    }

# '/' (default URL) ke liye ek 'route' banayein
@app.route("/")
def hello_world():
    return "Hello World! Resume Matcher Backend is running."

# File Upload aur Data Extraction ke liye API route
@app.route("/upload", methods=["POST"])
def upload_resume():
    """Secure upload handler with robust PDF extraction and graceful errors.
    Tries PyPDF2 first, then OCR via pdf2image+pytesseract if needed and poppler is available.
    """ 
    # 1. File validation
    if 'resume_file' not in request.files:
        return jsonify(error="No file part in the request", message="No file uploaded"), 400

    file = request.files['resume_file']
    if not file or file.filename == '':
        return jsonify(error="No selected file", message="Please select a file"), 400

    file_extension = os.path.splitext(file.filename)[1].lower()

    tmp_file = None
    try:
        # Save to a temporary file to allow libraries that need a path
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
        tmp_file = tmp.name
        file.save(tmp_file)
        tmp.close()

        file_content = ""

        # TXT
        if file_extension == '.txt':
            with open(tmp_file, 'r', encoding='utf-8', errors='ignore') as f:
                file_content = f.read()
            logger.info("TXT processed, length=%d", len(file_content))

        # PDF: try PyPDF2 text extraction first
        elif file_extension == '.pdf' and PDF_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(tmp_file)
                texts = []
                for p in reader.pages:
                    texts.append(p.extract_text() or "")
                file_content = "\n".join(texts)
                logger.info("PyPDF2 extracted length=%d", len(file_content))
            except Exception as e:
                logger.exception("PyPDF2 read failed: %s", e)
                file_content = ""

            # If too short, attempt OCR if available
            if len(file_content.strip()) < 50:
                if not OCR_AVAILABLE:
                    return jsonify(error="No text extractable from PDF and OCR libs not installed",
                                   message="Install pdf2image+pytesseract+Pillow and Poppler (pdftoppm)"), 400

                # Verify poppler (pdftoppm) is available
                pdftoppm_path = shutil.which('pdftoppm')
                poppler_kwargs = {}
                if pdftoppm_path:
                    poppler_dir = os.path.dirname(pdftoppm_path)
                    poppler_kwargs['poppler_path'] = poppler_dir
                    logger.info("Found pdftoppm at %s", pdftoppm_path)
                else:
                    logger.warning("pdftoppm not found in PATH; OCR cannot proceed")
                    return jsonify(error="Poppler not found (pdftoppm)",
                                   message="Install poppler and add to PATH. See README."), 500

                try:
                    # use convert_from_path on the saved file
                    from pdf2image import convert_from_path
                    images = convert_from_path(tmp_file, dpi=200, **poppler_kwargs)
                    ocr_text = ''
                    for img in images:
                        ocr_text += pytesseract.image_to_string(img, lang='eng') + '\n'
                    file_content = ocr_text
                    logger.info("OCR extracted length=%d", len(file_content))
                except Exception as e:
                    logger.exception("OCR failed: %s", e)
                    return jsonify(error="OCR failed", message=str(e)), 500

            if len(file_content.strip()) < 50:
                return jsonify(error="No text extractable from PDF", message="Try a text-based PDF or a different file"), 400

        # DOCX
        elif file_extension == '.docx' and DOCX_AVAILABLE:
            try:
                doc = Document(tmp_file)
                file_content = '\n'.join([para.text for para in doc.paragraphs])
                logger.info("DOCX extracted length=%d", len(file_content))
            except Exception as e:
                logger.exception("DOCX extraction failed: %s", e)
                return jsonify(error="DOCX extraction failed", message=str(e)), 500

        else:
            return jsonify(error=f"Unsupported file type: {file_extension}",
                           message="Supported: .txt, .pdf, .docx"), 400

        # final check
        if not file_content.strip():
            return jsonify(error="File is empty or no text could be extracted", message="Upload a file with readable content."), 400

        # Parsing and matching logic (use module-level STOPWORDS fallback)
        try:
            stopset = set(stopwords.words('english'))
        except Exception:
            stopset = STOPWORDS

        hr_skills_text = request.form.get('hr_skills', '')
        hr_skills_list = re.findall(r'\b\w+\b', hr_skills_text.lower())
        cleaned_resume_text = re.sub(r'[^a-zA-Z\s]',' ', file_content.lower())
        resume_words_list = cleaned_resume_text.split()

        hr_skills_clean = [w for w in hr_skills_list if w not in stopset]
        resume_words_clean = [w for w in resume_words_list if w not in stopset]

        hr_skills_set = set(hr_skills_clean)
        matched_words = hr_skills_set.intersection(resume_words_clean)
        score = len(matched_words)

        final_skills_list = [{"skill": s, "matched": s in matched_words} for s in hr_skills_set]
        len_hr_skills = len(hr_skills_set)
        percentage_match = float((score / len_hr_skills) * 100) if len_hr_skills > 0 else 0.0

        return jsonify(
            message="Data cleaned and initial matching completed.",
            matched_skills_count=score,
            total_hr_skills=len_hr_skills,
            detailed_skills=final_skills_list,
            percentage_match=percentage_match
        ), 200

    except UnicodeDecodeError as e:
        logger.exception("Unicode decode error: %s", e)
        return jsonify(error=f"File encoding issue: {str(e)}", message="Try saving as UTF-8 TXT."), 400
    except Exception as e:
        logger.exception("Unexpected extraction error: %s", e)
        return jsonify(error="Failed to extract text from file", message=str(e)), 500
    finally:
        if tmp_file and os.path.exists(tmp_file):
            try:
                os.remove(tmp_file)
            except Exception:
                logger.warning("Could not remove temp file %s", tmp_file)


def _extract_text_from_saved(tmp_file, file_extension):
    """Helper to extract text from a saved temp file. Returns extracted text (may be empty)."""
    file_content = ""
    # TXT
    if file_extension == '.txt':
        try:
            with open(tmp_file, 'r', encoding='utf-8', errors='ignore') as f:
                file_content = f.read()
            logger.info("TXT processed (batch), length=%d", len(file_content))
        except Exception:
            file_content = ""

    # PDF: try PyPDF2 text extraction first
    elif file_extension == '.pdf' and PDF_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(tmp_file)
            texts = []
            for p in reader.pages:
                texts.append(p.extract_text() or "")
            file_content = "\n".join(texts)
            logger.info("PyPDF2 extracted (batch) length=%d", len(file_content))
        except Exception:
            file_content = ""

        # If too short, attempt OCR if available
        if len(file_content.strip()) < 50:
            if OCR_AVAILABLE:
                pdftoppm_path = shutil.which('pdftoppm')
                poppler_kwargs = {}
                if pdftoppm_path:
                    poppler_dir = os.path.dirname(pdftoppm_path)
                    poppler_kwargs['poppler_path'] = poppler_dir
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(tmp_file, dpi=200, **poppler_kwargs)
                    ocr_text = ''
                    for img in images:
                        ocr_text += pytesseract.image_to_string(img, lang='eng') + '\n'
                    file_content = ocr_text
                    logger.info("OCR extracted (batch) length=%d", len(file_content))
                except Exception as e:
                    logger.exception("OCR failed (batch): %s", e)

    # DOCX
    elif file_extension == '.docx' and DOCX_AVAILABLE:
        try:
            doc = Document(tmp_file)
            file_content = '\n'.join([para.text for para in doc.paragraphs])
            logger.info("DOCX extracted (batch) length=%d", len(file_content))
        except Exception:
            file_content = ""

    return file_content


def _extract_tokens(text, stopset):
    cleaned = re.sub(r'[^a-zA-Z\s]', ' ', (text or '').lower())
    words = [w for w in cleaned.split() if len(w) > 2 and w not in stopset]
    return words


def _cosine_similarity_from_counters(counters):
    """Compute pairwise cosine similarity from a list of Counters (term frequencies)."""
    # Build vocabulary
    vocab = {}
    for c in counters:
        for k in c:
            if k not in vocab:
                vocab[k] = len(vocab)

    vecs = []
    for c in counters:
        v = [0.0] * len(vocab)
        for k, cnt in c.items():
            v[vocab[k]] = float(cnt)
        vecs.append(v)

    def dot(a, b):
        return sum(x*y for x, y in zip(a, b))

    def norm(a):
        return math.sqrt(sum(x*x for x in a))

    n = len(vecs)
    mat = [[0.0]*n for _ in range(n)]
    norms = [norm(v) for v in vecs]
    for i in range(n):
        for j in range(n):
            if norms[i] == 0 or norms[j] == 0:
                mat[i][j] = 0.0
            else:
                mat[i][j] = dot(vecs[i], vecs[j]) / (norms[i]*norms[j])
    return mat


@app.route('/batch-analyze', methods=['POST'])
def batch_analyze():
    """Accept up to 5 files (files[]) and return parsed outputs + pairwise similarity matrix.
    This duplicates extraction logic from /upload but keeps the single-file endpoint untouched.
    """
    files = request.files.getlist('files[]') or []
    # Backwards compatibility: allow single file input name
    if not files:
        single = request.files.get('resume_file')
        if single:
            files = [single]

    if not files:
        return jsonify(error='No files provided', message='Please upload one or more resume files using files[]'), 400

    if len(files) > 5:
        return jsonify(error='Too many files', message='Max 5 files allowed per batch'), 400

    # Prepare stopset
    try:
        stopset = set(stopwords.words('english'))
    except Exception:
        stopset = STOPWORDS

    results = [None] * len(files)

    def _process(idx, file_storage):
        tmp_path = None
        try:
            filename = file_storage.filename
            ext = os.path.splitext(filename)[1].lower()
            tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmp_path = tmpf.name
            file_storage.save(tmp_path)
            tmpf.close()
            text = _extract_text_from_saved(tmp_path, ext)
            tokens = _extract_tokens(text, stopset)
            ctr = Counter(tokens)
            snippet = (text or '').strip()[:800]
            return {'candidate': filename, 'text_snippet': snippet, 'tokens': list(ctr.keys()), 'counter': ctr}
        finally:
            try:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass

    # Run in parallel threads
    futures = []
    with ThreadPoolExecutor(max_workers=min(5, len(files))) as ex:
        for idx, f in enumerate(files):
            futures.append(ex.submit(_process, idx, f))

        parsed = []
        for fut in as_completed(futures):
            try:
                parsed.append(fut.result())
            except Exception as e:
                logger.exception('Error processing file in batch: %s', e)

    # Reorder parsed results to match uploaded files order when possible
    parsed_map = {p['candidate']: p for p in parsed}
    ordered = []
    remaining = [p for p in parsed]
    for f in files:
        if f.filename in parsed_map:
            ordered.append(parsed_map[f.filename])
            # also remove from remaining
            try:
                remaining.remove(parsed_map[f.filename])
            except Exception:
                pass
        else:
            if remaining:
                ordered.append(remaining.pop(0))

    if not ordered:
        ordered = parsed

    # Build counters list and compute similarity matrix
    counters = [p['counter'] for p in ordered]
    sim_matrix = _cosine_similarity_from_counters(counters) if counters else []

    # common skills intersection
    try:
        common_skills = set(ordered[0]['tokens']) if ordered else set()
        for p in ordered[1:]:
            common_skills = common_skills.intersection(set(p['tokens']))
    except Exception:
        common_skills = set()

    # HR required skills matching
    hr_skills_text = (request.form.get('hr_skills') or '').strip()
    required_set = set()
    if hr_skills_text:
        # tokenize simple comma/word list
        tokens = re.findall(r"\b[\w+#.+-]+\b", hr_skills_text.lower())
        try:
            stopset = set(stopwords.words('english'))
        except Exception:
            stopset = STOPWORDS
        required_set = set([t for t in tokens if t not in stopset and len(t) > 1])
    else:
        # optional: load skills.json if present
        skills_json_path = os.path.join(os.path.dirname(__file__), 'skills.json')
        if os.path.exists(skills_json_path):
            try:
                import json
                with open(skills_json_path, 'r', encoding='utf-8') as sf:
                    skills_data = json.load(sf)
                    if isinstance(skills_data, list):
                        required_set = set([s.lower() for s in skills_data if isinstance(s, str)])
            except Exception:
                required_set = set()

    # Prepare response (include matched skills and match pct)
    parsed_out = []
    for p in ordered:
        p_tokens = set(p.get('tokens', []))
        matched = sorted(list(required_set.intersection(p_tokens))) if required_set else []
        match_pct = (len(matched) / len(required_set) * 100.0) if required_set else 0.0
        selected = True if match_pct >= 80.0 else False
        parsed_out.append({
            'candidate': p['candidate'],
            'text_snippet': p['text_snippet'],
            'top_tokens': p['tokens'][:30],
            'required_skills': sorted(list(required_set)),
            'matched_skills': matched,
            'match_pct': round(match_pct, 2),
            'selected': selected
        })

    return jsonify({
        'parsed': parsed_out,
        'comparisons': {
            'similarity_matrix': sim_matrix,
            'common_skills': list(common_skills)
        }
    }), 200
 
# Yeh line check karti hai ki script ko direct run kiya gaya hai
if __name__ == "__main__":
    # Run without the reloader/debugger to avoid the double-process behavior
    app.run(host='127.0.0.1', port=5000, debug=False)
